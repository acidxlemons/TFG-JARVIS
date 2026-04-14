import os
import logging
import base64
import tempfile
import time
import uuid
import json
import subprocess
import re
import unicodedata
import hashlib
import mimetypes
from pathlib import Path
from datetime import datetime
from typing import List, Optional, Dict, Set

from fastapi import APIRouter, HTTPException, BackgroundTasks, Header, Body, Path as APIPath
from pydantic import BaseModel

from qdrant_client.models import (
    PointStruct,
    Filter,
    FieldCondition,
    MatchValue,
    VectorParams,
    Distance,
    SparseVectorParams,
)

try:
    from qdrant_client.models import Modifier
except ImportError:
    Modifier = None  # type: ignore[assignment]

from app.core.state import app_state
from app.core.permissions import can_write_without_tenant_header, get_all_collection_names, resolve_authorized_collections
from app.processing.embeddings.sentence_transformer import get_embedder, get_sparse_embedder, embed_texts

logger = logging.getLogger(__name__)

router = APIRouter(prefix="/documents", tags=["Documents"])


def _build_sparse_vector_params() -> SparseVectorParams:
    if Modifier is not None:
        return SparseVectorParams(modifier=Modifier.IDF)
    return SparseVectorParams()


def _normalize_search_text(text: str) -> str:
    normalized = unicodedata.normalize("NFKD", (text or "").strip().lower())
    return "".join(ch for ch in normalized if not unicodedata.combining(ch))


def _has_exact_id(text: str, token: str) -> bool:
    if not text or not token:
        return False
    pattern = rf"(?<![a-z0-9]){re.escape(token.lower())}(?![a-z0-9])"
    return re.search(pattern, _normalize_search_text(text)) is not None


def _document_type(collection_name: str, filename: str) -> str:
    coll = (collection_name or "").strip().lower()
    if coll.startswith("webs"):
        return "web"
    suffix = Path(filename or "").suffix.lower()
    if suffix in {".pdf", ".doc", ".docx", ".txt", ".md", ".log", ".json", ".xlsx", ".xls", ".csv", ".ppt", ".pptx"}:
        return "document"
    return "document"


def _compute_file_hash(file_path: str) -> str:
    sha256 = hashlib.sha256()
    with open(file_path, "rb") as fh:
        for chunk in iter(lambda: fh.read(1024 * 1024), b""):
            sha256.update(chunk)
    return sha256.hexdigest()


def _detect_source_type(metadata: Dict, collection_name: str, filename: str) -> str:
    source = str(metadata.get("source") or "").strip().lower()
    if source == "sharepoint":
        return "sharepoint"
    if source in {"scrape", "crawler"}:
        return "scrape"
    if source in {"web", "web_search"} or (collection_name or "").lower().startswith("webs"):
        return "web"
    return "upload"


def _build_registry_metadata(metadata: Dict, collection_name: str, tenant_id: Optional[str], filename: str) -> Dict:
    payload = dict(metadata or {})
    payload["collection_name"] = collection_name
    payload["tenant_id"] = tenant_id
    payload["active_collections"] = [collection_name] if collection_name else []
    payload.setdefault("source_path", str(payload.get("source_path") or payload.get("local_path") or filename))
    return payload


def _collect_documents_metadata(tenant_collections: List[str]) -> Dict[str, Dict[str, object]]:
    all_filenames: Dict[str, Dict[str, object]] = {}
    scroll_limit = max(100, int(os.getenv("DOCUMENTS_SCROLL_PAGE_SIZE", "1000")))
    max_points = max(scroll_limit, int(os.getenv("DOCUMENTS_METADATA_SCAN_LIMIT", "50000")))

    for collection_name in tenant_collections:
        try:
            offset = None
            scanned = 0
            while scanned < max_points:
                points, next_offset = app_state.qdrant.scroll(
                    collection_name=collection_name,
                    limit=scroll_limit,
                    offset=offset,
                    with_payload=True,
                    with_vectors=False,
                )
                if not points:
                    break

                for point in points:
                    payload = point.payload or {}
                    fname = str(payload.get("filename") or payload.get("source") or "unknown").strip()
                    if not fname:
                        continue
                    if fname not in all_filenames:
                        all_filenames[fname] = {
                            "filename": fname,
                            "collection": collection_name,
                            "chunks": 0,
                            "from_ocr": payload.get("from_ocr", False),
                            "ingested_at": payload.get("ingested_at", ""),
                            "type": _document_type(collection_name, fname),
                        }
                    all_filenames[fname]["chunks"] += 1

                scanned += len(points)
                if not next_offset:
                    break
                offset = next_offset

            if scanned >= max_points:
                logger.warning(
                    "Metadata scan truncado en coleccion %s tras %s puntos; ajuste DOCUMENTS_METADATA_SCAN_LIMIT si procede",
                    collection_name,
                    scanned,
                )
        except Exception as e:
            logger.warning(f"Error listando colección {collection_name}: {e}")

    return all_filenames


def _document_search_score(query: str, filename: str, collection_name: str) -> float:
    query_norm = _normalize_search_text(query)
    if not query_norm:
        return 0.0

    filename_norm = _normalize_search_text(filename)
    collection_norm = _normalize_search_text(collection_name)
    combined = f"{filename_norm} {collection_norm}"
    score = 0.0

    if query_norm == filename_norm:
        score += 120.0
    if query_norm in filename_norm:
        score += 80.0
    elif query_norm in combined:
        score += 40.0

    query_ids = re.findall(r"[a-z]{1,8}-\d{2,8}", query_norm)
    for token in query_ids:
        if _has_exact_id(filename, token):
            score += 40.0

    tokens = [token for token in re.findall(r"[a-z0-9]{2,}", query_norm) if token not in {"de", "la", "el", "los", "las", "del", "en"}]
    if tokens:
        filename_hits = sum(1 for token in tokens if token in filename_norm)
        collection_hits = sum(1 for token in tokens if token in collection_norm)
        score += filename_hits * 8.0
        score += collection_hits * 4.0
        if filename_hits == len(tokens):
            score += 20.0
        elif filename_hits >= max(1, len(tokens) - 1):
            score += 8.0

    return score

# ======================================================
# BATCH UPLOAD (BASE64) AND BACKGROUND PROCESSING
# ======================================================


def _is_hybrid_collection(collection_name: str) -> bool:
    try:
        info = app_state.qdrant.get_collection(collection_name)
        vectors = getattr(info.config.params, "vectors", None)
        sparse_vectors_cfg = getattr(info.config.params, "sparse_vectors", None)
        has_named_dense = isinstance(vectors, dict) and "dense" in vectors
        has_sparse = bool(sparse_vectors_cfg)
        return has_named_dense and has_sparse
    except Exception:
        return False


def _ensure_hybrid_collection(collection_name: str, vector_dim: Optional[int] = None) -> Dict[str, object]:
    if not isinstance(collection_name, str) or not collection_name.strip():
        raise ValueError("collection_name inválida")
    collection_name = collection_name.strip()

    try:
        collections = app_state.qdrant.get_collections().collections
        exists = any(c.name == collection_name for c in collections)
    except Exception as e:
        raise RuntimeError(f"No se pudo consultar colecciones Qdrant: {e}") from e

    if not exists:
        if vector_dim is None:
            vector_dim = get_embedder().dimension
        app_state.qdrant.create_collection(
            collection_name=collection_name,
            vectors_config={"dense": VectorParams(size=int(vector_dim), distance=Distance.COSINE)},
            sparse_vectors_config={"sparse": _build_sparse_vector_params()},
        )
        return {
            "collection": collection_name,
            "created": True,
            "exists": True,
            "hybrid": True,
        }

    hybrid_ok = _is_hybrid_collection(collection_name)
    return {
        "collection": collection_name,
        "created": False,
        "exists": True,
        "hybrid": hybrid_ok,
    }


def _extract_native_text(file_path: str, suffix: str) -> str:
    """Extrae texto nativo según extensión (sin OCR)."""
    suffix = (suffix or "").lower()

    if suffix == ".pdf":
        from app.processing.ocr.paddle_ocr import extract_text_native
        return extract_text_native(file_path) or ""

    if suffix == ".docx":
        try:
            from docx import Document as DocxDocument

            doc = DocxDocument(file_path)
            lines: List[str] = []

            for p in doc.paragraphs:
                txt = (p.text or "").strip()
                if txt:
                    lines.append(txt)

            for table in doc.tables:
                for row in table.rows:
                    cells = [((c.text or "").strip()) for c in row.cells]
                    cells = [c for c in cells if c]
                    if cells:
                        lines.append(" | ".join(cells))

            return "\n".join(lines)
        except Exception as e:
            logger.warning(f"No se pudo extraer DOCX {Path(file_path).name}: {e}")
            return ""

    if suffix == ".doc":
        # Formato binario legacy: intentamos antiword si está disponible.
        try:
            proc = subprocess.run(
                ["antiword", file_path],
                capture_output=True,
                text=True,
                timeout=30,
                check=False,
            )
            if proc.returncode == 0:
                return (proc.stdout or "").strip()
        except Exception:
            pass
        return ""

    if suffix in {".xlsx", ".xls"}:
        try:
            import pandas as pd

            workbook = pd.ExcelFile(file_path)
            chunks: List[str] = []
            for sheet_name in workbook.sheet_names:
                try:
                    df = pd.read_excel(workbook, sheet_name=sheet_name, dtype=str, header=None)
                except Exception:
                    continue

                df = df.fillna("")
                rows: List[str] = []
                for row in df.values.tolist():
                    cols = [str(c).strip() for c in row if str(c).strip()]
                    if cols:
                        rows.append(" | ".join(cols))

                if rows:
                    chunks.append(f"=== SHEET: {sheet_name} ===\n" + "\n".join(rows))

            return "\n\n".join(chunks)
        except Exception as e:
            logger.warning(f"No se pudo extraer Excel {Path(file_path).name}: {e}")
            return ""

    if suffix == ".csv":
        try:
            import pandas as pd

            df = pd.read_csv(file_path, dtype=str, header=None, keep_default_na=False)
            rows: List[str] = []
            for row in df.values.tolist():
                cols = [str(c).strip() for c in row if str(c).strip()]
                if cols:
                    rows.append(" | ".join(cols))
            return "\n".join(rows)
        except Exception as e:
            logger.warning(f"No se pudo extraer CSV {Path(file_path).name}: {e}")
            return ""

    return ""

def process_document(
    file_path: str,
    filename: str,
    metadata: dict,
    tenant_id: Optional[str],
    target_collection: Optional[str] = None,
):
    """Procesa documento multiplataforma (TXT, PDF, Word, Excel, etc.) en background"""
    logger.info(f"Procesando documento: {filename}")
    app_state.memory.update_ingestion_status(filename, "processing", "Iniciando procesamiento...")

    file_hash = ""
    file_size = None
    mime_type = None
    collection_name = target_collection or os.getenv("QDRANT_COLLECTION", "documents")
    registry_metadata = _build_registry_metadata(metadata or {}, collection_name, tenant_id, filename)
    source_type = _detect_source_type(registry_metadata, collection_name, filename)
    source_path = str(registry_metadata.get("source_path") or filename)

    try:
        file_hash = _compute_file_hash(file_path)
        file_size = Path(file_path).stat().st_size
        mime_type = mimetypes.guess_type(filename)[0]
        suffix = Path(file_path).suffix.lower()

        # 1) Texto plano
        if suffix in {".txt", ".md", ".log", ".json"}:
            logger.info(f"Leyendo texto plano de {filename}...")
            try:
                with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
                    text = f.read()
                from_ocr = False
                pages = 1
            except Exception as e:
                logger.error(f"Error leyendo texto plano de {filename}: {e}")
                text = ""
                from_ocr = False
                pages = 0
        elif suffix in {".jpg", ".jpeg", ".png", ".tiff", ".bmp", ".webp"}:
            logger.info(f"Aplicando OCR a imagen {filename}...")
            result = app_state.ocr_pipeline.process_file(file_path)
            text = result.text
            pages = max(1, result.pages)
            from_ocr = True
        elif suffix == ".pdf":
            # PDFs: decide OCR din?micamente (escaneado vs nativo).
            needs_ocr = app_state.ocr_pipeline.needs_ocr(file_path)
            if needs_ocr:
                logger.info(f"Aplicando OCR a {filename}...")
                result = app_state.ocr_pipeline.process_file(file_path)
                text = result.text
                pages = result.pages
                from_ocr = True
            else:
                logger.info(f"Extrayendo texto nativo de PDF {filename}...")
                text = _extract_native_text(file_path, suffix)
                pages = max(1, text.count("=== P?GINA ===")) if text else 0
                from_ocr = False

            # Fallback defensivo: si PDF nativo devuelve texto muy corto, intentar OCR.
            if not from_ocr and len((text or "").strip()) < 50:
                try:
                    logger.info(f"Fallback OCR por texto corto en {filename}...")
                    result = app_state.ocr_pipeline.process_file(file_path)
                    if len((result.text or "").strip()) > len((text or "").strip()):
                        text = result.text
                        pages = result.pages
                        from_ocr = True
                except Exception as e:
                    logger.warning(f"Fallback OCR fall? para {filename}: {e}")
        else:
            # 2) Office/otros formatos con parser nativo.
            logger.info(f"Extrayendo texto nativo de {filename} ({suffix})...")
            text = _extract_native_text(file_path, suffix)
            pages = max(1, text.count("=== P?GINA ===")) if text else 0
            from_ocr = False

        if not text or len(text.strip()) < 50:
            logger.warning(f"Texto insuficiente en {filename}")
            app_state.memory.update_ingestion_status(
                filename,
                "failed",
                "Texto insuficiente para indexar",
            )
            app_state.memory.upsert_document_record(
                filename=filename,
                source_path=source_path,
                source_type=source_type,
                file_hash=file_hash,
                file_size=file_size,
                mime_type=mime_type,
                page_count=pages,
                chunk_count=0,
                from_ocr=from_ocr,
                metadata=registry_metadata,
                status="failed",
            )
            app_state.memory.log_audit_event(
                action="document_index_failed",
                resource_type="document",
                details={
                    "filename": filename,
                    "collection_name": collection_name,
                    "reason": "Texto insuficiente para indexar",
                    "source_type": source_type,
                },
            )
            Path(file_path).unlink(missing_ok=True)
            return

        # 3) Chunking (usando chunker robusto del estado global)
        chunks = app_state.chunker.chunk_text(text, filename)
        logger.info(f"Generados {len(chunks)} chunks")

        # 4) Vectorizaci?n (Dense + Sparse para B?squeda H?brida)
        texts = [c["text"] for c in chunks]
        embedder = get_embedder()

        hybrid_ok = True
        try:
            ensured = _ensure_hybrid_collection(collection_name, vector_dim=embedder.dimension)
            hybrid_ok = bool(ensured.get("hybrid"))
            if ensured.get("created"):
                logger.info("Colecci?n '%s' creada autom?ticamente (h?brida).", collection_name)
            if not hybrid_ok:
                logger.warning("Colecci?n '%s' sin h?brido. Indexando dense-only.", collection_name)
        except Exception as e:
            hybrid_ok = False
            logger.warning("Error verificando colecci?n Qdrant '%s': %s. Indexando dense-only.", collection_name, e)

        logger.info(f"Generando embeddings para {len(texts)} chunks...")
        dense_vectors = embedder.encode_batch(texts)

        if hybrid_ok:
            sparse_embedder = get_sparse_embedder()
            if sparse_embedder is None:
                logger.warning("Sparse embedder no disponible. Indexando '%s' en dense-only.", collection_name)
                sparse_vectors = [None] * len(texts)
                hybrid_ok = False
            else:
                sparse_gen = list(sparse_embedder.embed(texts))
                sparse_vectors = [{"indices": v.indices.tolist(), "values": v.values.tolist()} for v in sparse_gen]
        else:
            sparse_vectors = [None] * len(texts)

        now = datetime.utcnow()
        now_iso = now.isoformat() + "Z"
        now_ts = int(now.timestamp())

        # 5) Insertar en Qdrant
        points = []
        for chunk, d_vec, s_vec in zip(chunks, dense_vectors, sparse_vectors):
            payload = {
                "text": chunk["text"],
                "filename": filename,
                "source": metadata.get("source") or filename,
                "page": chunk["metadata"].get("page"),
                "chunk_index": chunk["metadata"].get("chunk_index"),
                "from_ocr": from_ocr,
                "ingested_at": now_iso,
                "ingested_at_ts": now_ts,
                "tenant_id": tenant_id,
                **metadata,
            }
            if s_vec is None:
                vector_dict = d_vec
            else:
                vector_dict = {
                    "dense": d_vec,
                    "sparse": s_vec
                }
            points.append(
                PointStruct(
                    id=str(uuid.uuid4()),
                    vector=vector_dict,
                    payload=payload,
                )
            )

        app_state.qdrant.upsert(collection_name=collection_name, points=points)

        logger.info(f"? Documento procesado e indexado: {filename} ({len(points)} chunks)")
        app_state.memory.update_ingestion_status(
            filename,
            "completed",
            f"Indexado correctamente ({len(points)} chunks)",
        )
        document_row = app_state.memory.upsert_document_record(
            filename=filename,
            source_path=source_path,
            source_type=source_type,
            file_hash=file_hash,
            file_size=file_size,
            mime_type=mime_type,
            page_count=pages,
            chunk_count=len(points),
            from_ocr=from_ocr,
            metadata=registry_metadata,
            status="indexed",
        )
        app_state.memory.log_audit_event(
            action="document_indexed",
            resource_type="document",
            resource_id=document_row["id"],
            details={
                "filename": filename,
                "collection_name": collection_name,
                "chunk_count": len(points),
                "source_type": source_type,
                "from_ocr": from_ocr,
            },
        )

    except Exception as e:
        logger.error(f"Error procesando documento {filename}: {e}")
        app_state.memory.update_ingestion_status(filename, "failed", str(e))
        if file_hash:
            try:
                app_state.memory.upsert_document_record(
                    filename=filename,
                    source_path=source_path,
                    source_type=source_type,
                    file_hash=file_hash,
                    file_size=file_size,
                    mime_type=mime_type,
                    page_count=0,
                    chunk_count=0,
                    from_ocr=False,
                    metadata=registry_metadata,
                    status="failed",
                )
                app_state.memory.log_audit_event(
                    action="document_index_failed",
                    resource_type="document",
                    details={
                        "filename": filename,
                        "collection_name": collection_name,
                        "source_type": source_type,
                        "error": str(e),
                    },
                )
            except Exception as audit_error:
                logger.warning("No se pudo registrar fallo de documento %s: %s", filename, audit_error)
    finally:
        Path(file_path).unlink(missing_ok=True)


@router.post("/upload")
async def upload_document(
    background_tasks: BackgroundTasks,
    payload: dict = Body(...),
    x_tenant_id: Optional[str] = Header(None, alias="X-Tenant-Id"),
):
    """Sube y procesa documento mediante Base64 para el backend monolítico."""
    try:
        filename = payload.get("filename")
        content_base64 = payload.get("content_base64")
        metadata = payload.get("metadata") or {}

        if not isinstance(filename, str) or not filename.strip():
            raise HTTPException(400, "Campo 'filename' es requerido")
        if not isinstance(content_base64, str) or not content_base64.strip():
            raise HTTPException(400, "Campo 'content_base64' es requerido")

        content = base64.b64decode(content_base64)

        with tempfile.NamedTemporaryFile(delete=False, suffix=Path(filename).suffix) as tmp:
            tmp.write(content)
            tmp_path = tmp.name

        default_tenant = os.getenv("DEFAULT_TENANT_ID")
        if not x_tenant_id and not can_write_without_tenant_header():
            raise HTTPException(403, "X-Tenant-Id es obligatorio para indexar documentos")

        # Multi-site: indexer passes destination collection in X-Tenant-Id.
        # Backward compatibility: if header is default tenant id, use default collection.
        target_collection = (x_tenant_id or "").strip()
        if target_collection == default_tenant:
            target_collection = os.getenv("QDRANT_COLLECTION", "documents")
        if not target_collection:
            target_collection = os.getenv("QDRANT_COLLECTION", "documents")

        background_tasks.add_task(
            process_document,
            file_path=tmp_path,
            filename=filename,
            metadata=metadata,
            tenant_id=x_tenant_id or os.getenv("DEFAULT_TENANT_ID"),
            target_collection=target_collection,
        )

        return {
            "status": "processing",
            "filename": filename,
            "collection": target_collection,
            "message": "Documento en cola",
        }
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(500, detail=str(e))


class IngestionStatusUpdate(BaseModel):
    filename: str
    status: str
    message: Optional[str] = None


class SharePointSyncReport(BaseModel):
    site_id: str
    folder_path: str
    collection_name: Optional[str] = None
    site_name: Optional[str] = None
    delta_token: Optional[str] = None
    last_sync: Optional[str] = None
    status: str = "success"
    downloaded_files: int = 0
    indexed_files: int = 0
    errors: int = 0
    message: Optional[str] = None
    is_active: bool = True
    subscription_id: Optional[str] = None
    subscription_expires: Optional[str] = None


@router.post("/status")
async def update_document_status(payload: IngestionStatusUpdate):
    """Actualiza estado de ingesta (compatibilidad con indexer)."""
    try:
        if not payload.filename.strip():
            raise HTTPException(400, detail="filename requerido")
        if not payload.status.strip():
            raise HTTPException(400, detail="status requerido")

        app_state.memory.update_ingestion_status(
            payload.filename.strip(),
            payload.status.strip(),
            payload.message,
        )
        return {"status": "ok"}
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(500, detail=str(e))


@router.delete("/delete")
async def delete_document_by_filename(
    filename: str,
    x_tenant_id: Optional[str] = Header(None, alias="X-Tenant-Id"),
):
    """Elimina por filename en una colección concreta."""
    try:
        if not isinstance(filename, str) or not filename.strip():
            raise HTTPException(400, detail="filename requerido")
        if not x_tenant_id and not can_write_without_tenant_header():
            raise HTTPException(403, detail="X-Tenant-Id es obligatorio para eliminar documentos")

        collection_name = x_tenant_id or os.getenv("QDRANT_COLLECTION", "documents")
        target_name = filename.strip()

        # Qdrant no elimina por payload directamente en todos los clientes/versiones:
        # primero buscamos ids y luego borramos por ids.
        ids_to_delete = []
        offset = None

        while True:
            points, next_offset = app_state.qdrant.scroll(
                collection_name=collection_name,
                scroll_filter=Filter(
                    must=[FieldCondition(key="filename", match=MatchValue(value=target_name))]
                ),
                limit=500,
                offset=offset,
                with_payload=False,
                with_vectors=False,
            )
            if not points:
                break

            ids_to_delete.extend([p.id for p in points if getattr(p, "id", None) is not None])
            if not next_offset:
                break
            offset = next_offset

        deleted = len(ids_to_delete)
        if deleted > 0:
            app_state.qdrant.delete(
                collection_name=collection_name,
                points_selector=ids_to_delete,
            )

        app_state.memory.update_ingestion_status(
            target_name,
            "deleted",
            f"Eliminado de {collection_name} ({deleted} chunks)",
        )
        updated_rows = app_state.memory.mark_document_deleted(
            target_name,
            collection_name=collection_name,
        )
        app_state.memory.log_audit_event(
            action="document_deleted",
            resource_type="document",
            details={
                "filename": target_name,
                "collection_name": collection_name,
                "deleted_chunks": deleted,
                "registry_rows_updated": updated_rows,
            },
        )

        return {
            "status": "ok",
            "collection": collection_name,
            "filename": target_name,
            "deleted_chunks": deleted,
        }
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(500, detail=str(e))

# ======================================================
# MONITORING AND MANAGEMENT
# ======================================================

@router.get("/stats")
async def get_documents_stats():
    """Estadísticas de la colección"""
    try:
        collection_name = os.getenv("QDRANT_COLLECTION", "documents")
        res = app_state.qdrant.count(collection_name=collection_name, exact=True)
        points_count = res.count if hasattr(res, "count") else res.get("count", 0)
        return {"collection_name": collection_name, "status": "ok", "points_count": points_count}
    except Exception as e:
        message = str(e)
        if "doesn't exist" in message or "Not found: Collection" in message or "404" in message:
            return {"collection_name": collection_name, "status": "missing", "points_count": 0}
        raise HTTPException(500, detail=message)


@router.get("/list")
async def list_documents(
    x_tenant_id: Optional[str] = Header(None, alias="X-Tenant-Id"),
    x_tenant_ids: Optional[str] = Header(None, alias="X-Tenant-Ids"),
):
    """Lista documentos indexados extrayendo filenames del vector store"""
    try:
        tenant_collections = resolve_authorized_collections(
            qdrant_client=app_state.qdrant,
            x_tenant_id=x_tenant_id,
            x_tenant_ids=x_tenant_ids,
        )

        if not tenant_collections:
            return {
                "status": "ok",
                "total": 0,
                "collections_searched": [],
                "documents": [],
            }

        all_filenames = _collect_documents_metadata(tenant_collections)

        return {
            "status": "ok",
            "total": len(all_filenames),
            "collections_searched": tenant_collections,
            "documents": sorted(all_filenames.values(), key=lambda d: d["filename"]),
        }
    except Exception as e:
        raise HTTPException(500, detail=str(e))


@router.get("/search")
async def search_documents(
    q: str,
    limit: int = 10,
    x_tenant_id: Optional[str] = Header(None, alias="X-Tenant-Id"),
    x_tenant_ids: Optional[str] = Header(None, alias="X-Tenant-Ids"),
):
    """Busca documentos por nombre/código respetando las colecciones autorizadas."""
    try:
        query = (q or "").strip()
        if len(query) < 2:
            raise HTTPException(400, detail="Query demasiado corta")

        limit = max(1, min(int(limit), 100))
        tenant_collections = resolve_authorized_collections(
            qdrant_client=app_state.qdrant,
            x_tenant_id=x_tenant_id,
            x_tenant_ids=x_tenant_ids,
        )
        if not tenant_collections:
            return {
                "status": "ok",
                "query": query,
                "total": 0,
                "shown": 0,
                "collections_searched": [],
                "results": [],
            }

        documents = _collect_documents_metadata(tenant_collections)
        matches = []
        for document in documents.values():
            score = _document_search_score(
                query=query,
                filename=str(document.get("filename", "")),
                collection_name=str(document.get("collection", "")),
            )
            if score <= 0:
                continue
            row = dict(document)
            row["score"] = round(score, 2)
            matches.append(row)

        matches.sort(
            key=lambda d: (
                float(d.get("score", 0.0)),
                int(d.get("chunks", 0)),
                str(d.get("filename", "")),
            ),
            reverse=True,
        )
        return {
            "status": "ok",
            "query": query,
            "total": len(matches),
            "shown": min(len(matches), limit),
            "collections_searched": tenant_collections,
            "results": matches[:limit],
        }
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(500, detail=str(e))


@router.get("/ingestion-status")
async def get_ingestion_status(limit: int = 20):
    """Consulta la memoria de Postgres para ver el estado de ingestión de archivos"""
    try:
        return app_state.memory.get_ingestion_status(limit=limit)
    except Exception as e:
        raise HTTPException(500, detail=str(e))


@router.get("/registry")
async def get_document_registry(
    limit: int = 100,
    status: Optional[str] = None,
    source_type: Optional[str] = None,
):
    """Consulta el registro SQL persistido de documentos indexados."""
    try:
        rows = app_state.memory.get_document_registry(
            limit=limit,
            status=status,
            source_type=source_type,
        )
        return {
            "status": "ok",
            "total": len(rows),
            "documents": rows,
        }
    except Exception as e:
        raise HTTPException(500, detail=str(e))


@router.post("/sharepoint-sync/report")
async def report_sharepoint_sync(payload: SharePointSyncReport):
    """Persistencia SQL del estado de sincronización SharePoint."""
    try:
        row_id = app_state.memory.update_sharepoint_sync_state(
            site_id=payload.site_id.strip(),
            folder_path=payload.folder_path.strip(),
            delta_token=payload.delta_token,
            last_sync=payload.last_sync,
            subscription_id=payload.subscription_id,
            subscription_expires=payload.subscription_expires,
            is_active=payload.is_active,
        )
        app_state.memory.log_audit_event(
            action=f"sharepoint_sync_{payload.status.strip().lower()}",
            resource_type="sharepoint_sync",
            resource_id=row_id,
            details={
                "site_id": payload.site_id,
                "site_name": payload.site_name,
                "folder_path": payload.folder_path,
                "collection_name": payload.collection_name,
                "downloaded_files": payload.downloaded_files,
                "indexed_files": payload.indexed_files,
                "errors": payload.errors,
                "message": payload.message,
                "delta_token_present": bool(payload.delta_token),
            },
        )
        return {"status": "ok", "id": row_id}
    except Exception as e:
        raise HTTPException(500, detail=str(e))


@router.get("/sharepoint-sync")
async def get_sharepoint_sync(limit: int = 50):
    """Lista el último estado SQL de sincronización por sitio SharePoint."""
    try:
        rows = app_state.memory.get_sharepoint_sync_states(limit=limit)
        return {
            "status": "ok",
            "total": len(rows),
            "items": rows,
        }
    except Exception as e:
        raise HTTPException(500, detail=str(e))


@router.get("/audit-log")
async def get_audit_log(
    limit: int = 100,
    action: Optional[str] = None,
    resource_type: Optional[str] = None,
):
    """Lista eventos de auditoría operativa recientes."""
    try:
        rows = app_state.memory.get_audit_log(
            limit=limit,
            action=action,
            resource_type=resource_type,
        )
        return {
            "status": "ok",
            "total": len(rows),
            "items": rows,
        }
    except Exception as e:
        raise HTTPException(500, detail=str(e))


def _collection_aliases(name: str) -> List[str]:
    aliases: List[str] = []
    if name.startswith("documents_"):
        suffix = name[len("documents_"):].lower()
        if suffix:
            aliases.append(f"docs_{suffix}")
    if name.startswith("docs_"):
        suffix = name[len("docs_"):].upper()
        if suffix:
            aliases.append(f"documents_{suffix}")
    return aliases


def _count_unique_docs(collection_name: str, max_points: int = 20000) -> Dict[str, int]:
    """Cuenta documentos únicos por filename/source con límite defensivo."""
    seen: Set[str] = set()
    offset = None
    scanned = 0

    while True:
        points, next_offset = app_state.qdrant.scroll(
            collection_name=collection_name,
            limit=500,
            offset=offset,
            with_payload=True,
            with_vectors=False,
        )
        if not points:
            break

        for point in points:
            payload = point.payload or {}
            fname = payload.get("filename") or payload.get("source")
            if isinstance(fname, str) and fname:
                seen.add(fname)

        scanned += len(points)
        if not next_offset or scanned >= max_points:
            break
        offset = next_offset

    return {
        "documents": len(seen),
        "scanned_points": scanned,
        "truncated": 1 if scanned >= max_points else 0,
    }


@router.get("/collections/status")
async def get_collections_status(collections: Optional[str] = None):
    """Estado real de colecciones: existencia + puntos + documentos."""
    try:
        existing = {c.name for c in app_state.qdrant.get_collections().collections if getattr(c, "name", None)}
        requested = [c.strip() for c in (collections or "").split(",") if c.strip()]
        if not requested:
            requested = sorted(existing)

        def _collection_counters(name: str) -> Dict[str, int]:
            cnt = app_state.qdrant.count(collection_name=name, exact=True)
            points = cnt.count if hasattr(cnt, "count") else cnt.get("count", 0)
            doc_stats = _count_unique_docs(name)
            return {
                "points": int(points),
                "documents": int(doc_stats["documents"]),
                "scanned_points": int(doc_stats["scanned_points"]),
                "truncated": int(doc_stats["truncated"]),
            }

        rows = []
        for requested_name in requested:
            resolved_name = requested_name
            exists = requested_name in existing
            alias_used = None

            if not exists:
                for alias in _collection_aliases(requested_name):
                    if alias in existing:
                        exists = True
                        resolved_name = alias
                        alias_used = alias
                        break

            points_count = 0
            documents_count = 0
            scanned_points = 0
            truncated = 0
            if exists:
                counters = _collection_counters(resolved_name)
                points_count = counters["points"]
                documents_count = counters["documents"]
                scanned_points = counters["scanned_points"]
                truncated = counters["truncated"]

            # Compatibilidad legacy: si la colección solicitada existe vacía pero su alias legacy tiene docs,
            # reflejarlo para que el status operativo en OpenWebUI no sea engañoso.
            fallback_alias = None
            fallback_documents = 0
            for alias in _collection_aliases(requested_name):
                if alias in existing:
                    try:
                        alias_docs = _collection_counters(alias)["documents"]
                        if alias_docs > fallback_documents:
                            fallback_documents = alias_docs
                            fallback_alias = alias
                    except Exception:
                        continue

            display_documents = max(documents_count, fallback_documents)

            rows.append(
                {
                    "requested_collection": requested_name,
                    "resolved_collection": resolved_name,
                    "exists": exists,
                    "alias_used": alias_used,
                    "points": points_count,
                    "documents": documents_count,
                    "display_documents": display_documents,
                    "fallback_alias": fallback_alias,
                    "fallback_documents": fallback_documents,
                    "scanned_points": scanned_points,
                    "truncated": bool(truncated),
                }
            )

        return {
            "status": "ok",
            "requested_total": len(requested),
            "existing_total": len(existing),
            "existing_collections": sorted(existing),
            "collections": rows,
        }
    except Exception as e:
        raise HTTPException(500, detail=str(e))


class EnsureCollectionsRequest(BaseModel):
    collections: List[str]


@router.post("/collections/ensure")
async def ensure_collections(payload: EnsureCollectionsRequest):
    """Crea colecciones híbridas (dense+sparse) si no existen."""
    try:
        requested = [c.strip() for c in payload.collections if isinstance(c, str) and c.strip()]
        if not requested:
            raise HTTPException(400, detail="Debe indicar al menos una colección válida")

        results = []
        errors = []
        for name in list(dict.fromkeys(requested)):
            try:
                result = _ensure_hybrid_collection(name)
                results.append(result)
            except Exception as e:
                logger.error("Error asegurando colección '%s': %s", name, e)
                errors.append({"collection": name, "error": str(e)})

        return {
            "status": "ok" if not errors else "partial",
            "requested": len(requested),
            "ensured": len(results),
            "created": len([r for r in results if r.get("created")]),
            "hybrid_ok": len([r for r in results if r.get("hybrid")]),
            "results": results,
            "errors": errors,
        }
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(500, detail=str(e))
